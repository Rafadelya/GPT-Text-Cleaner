import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pyperclip
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import re
import os
import io
import base64
from PIL import Image, ImageTk, ImageGrab
import pytesseract
from langdetect import detect, LangDetectException
from io import BytesIO
import sys
import subprocess

## so simple
class GPTTextCleaner:
    def __init__(self, root):
        self.root = root
        self.root.title("GPT Text Cleaner")
        self.root.geometry("800x600")
        
        # Set tesseract path
        tesseract_path = '/usr/local/bin/tesseract'
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Configure style for blue button
        style = ttk.Style()
        style.configure('TButton', foreground='white', background='#007bff')
        style.map('TButton',
                 foreground=[('pressed', 'white'), ('active', 'white')],
                 background=[('pressed', '#0056b3'), ('active', '#007bff'), ('!active', '#007bff')])
        
        # Configure colors
        self.root.configure(bg='white')
        
        # Detect OS and Tesseract status
        self.os_name = sys.platform
        self.tesseract_status = "Installed" if os.path.exists(self.get_tesseract_path()) else "Not Found"

        # Add a status bar
        self.status_bar = ttk.Label(self.root, text=f"OS: {self.os_name} | Tesseract: {self.tesseract_status}", 
                                  anchor="w", background='white')
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        notebook = ttk.Notebook(root)
        notebook.pack(fill=tk.BOTH, expand=True)

        # Text tab
        text_frame = ttk.Frame(notebook)
        notebook.add(text_frame, text="Text")

        # Screenshot tab
        screenshot_frame = ttk.Frame(notebook)
        notebook.add(screenshot_frame, text="Screenshot")

        # Setup tabs
        self.setup_text_tab(text_frame)
        self.setup_screenshot_tab(screenshot_frame)

    def get_tesseract_path(self):
        if self.os_name == 'darwin':
            return '/usr/local/bin/tesseract' if os.path.exists('/usr/local/bin/tesseract') else '/opt/homebrew/bin/tesseract'
        elif self.os_name == 'win32':
            return r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        else:
            return '/usr/bin/tesseract'

    def install_tesseract(self):
        if self.os_name == 'darwin':
            try:
                # Проверяем наличие Homebrew
                result = subprocess.run(['which', 'brew'], capture_output=True, text=True)
                if result.returncode != 0:
                    messagebox.showerror("Ошибка", "Homebrew не установлен. Пожалуйста, установите Homebrew с brew.sh")
                    return

                # Устанавливаем Tesseract и языковые пакеты
                subprocess.run(['brew', 'install', 'tesseract', 'tesseract-lang'], check=True)
                messagebox.showinfo("Успех", "Tesseract OCR успешно установлен")
                
                # Обновляем статус
                self.tesseract_status = "Установлен"
                self.tesseract_label.config(text=f"Tesseract: {self.tesseract_status}")
                self.status_bar.config(text=f"OS: {self.os_name} | Tesseract: {self.tesseract_status}")
            except subprocess.CalledProcessError as e:
                messagebox.showerror("Ошибка", f"Ошибка при установке Tesseract: {str(e)}")
        elif self.os_name == 'linux':
            try:
                # Устанавливаем Tesseract и языковые пакеты для русского и английского
                subprocess.run(['sudo', 'apt', 'update'], check=True)
                subprocess.run(['sudo', 'apt', 'install', '-y', 'tesseract-ocr', 'tesseract-ocr-rus', 'tesseract-ocr-eng'], check=True)
                messagebox.showinfo("Успех", "Tesseract OCR успешно установлен")
                
                # Обновляем статус
                self.tesseract_status = "Установлен"
                self.tesseract_label.config(text=f"Tesseract: {self.tesseract_status}")
                self.status_bar.config(text=f"OS: {self.os_name} | Tesseract: {self.tesseract_status}")
            except subprocess.CalledProcessError as e:
                messagebox.showerror("Ошибка", f"Ошибка при установке Tesseract: {str(e)}")
        elif self.os_name == 'win32':
            messagebox.showinfo("Информация",
                              "Для Windows требуется ручная установка Tesseract OCR:\n\n"
                              "1. Скачайте установщик с github.com/UB-Mannheim/tesseract/wiki\n"
                              "2. Запустите установщик и следуйте инструкциям\n"
                              "3. При установке отметьте дополнительные языки (русский)\n"
                              "4. Убедитесь, что путь установки: C:\\Program Files\\Tesseract-OCR")

    def setup_text_tab(self, parent):
        self.text_input = tk.Text(parent, height=12, width=80, bg='white', fg='black')
        self.text_input.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(pady=5, padx=10, fill=tk.X)

        self.btn_load_file = ttk.Button(btn_frame, text="Load File", command=self.load_file)
        self.btn_load_file.pack(side=tk.LEFT, padx=5)

        self.btn_load_clipboard = ttk.Button(btn_frame, text="From Clipboard", command=self.load_from_clipboard)
        self.btn_load_clipboard.pack(side=tk.LEFT, padx=5)

        self.btn_process = ttk.Button(btn_frame, text="Process", command=self.process_text)
        self.btn_process.pack(side=tk.RIGHT, padx=5)

        self.btn_copy_result = ttk.Button(btn_frame, text="Copy Result", command=self.copy_to_clipboard)
        self.btn_copy_result.pack(side=tk.RIGHT, padx=5)

        self.result_text = tk.Text(parent, height=12, width=80, bg='white', fg='black')
        self.result_text.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        save_frame = ttk.Frame(parent)
        save_frame.pack(pady=5, padx=10, fill=tk.X)

        self.save_btn = ttk.Button(save_frame, text="Save as Word", command=self.save_to_word, style='TButton')
        self.save_btn.pack(side=tk.LEFT, padx=5)
    
    def setup_screenshot_tab(self, parent):
        top_frame = ttk.Frame(parent)
        top_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.btn_load_image = ttk.Button(top_frame, text="Load Image", command=self.load_image)
        self.btn_load_image.pack(side=tk.LEFT, padx=5)
        
        self.btn_clipboard_image = ttk.Button(top_frame, text="Image from Clipboard", command=self.load_clipboard_image)
        self.btn_clipboard_image.pack(side=tk.LEFT, padx=5)
        
        self.btn_process_image = ttk.Button(top_frame, text="Recognize Text", command=self.process_image)
        self.btn_process_image.pack(side=tk.RIGHT, padx=5)
        
        # Image frame
        self.image_frame = ttk.LabelFrame(parent, text="Image")
        self.image_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.image_label = ttk.Label(self.image_frame)
        self.image_label.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # OCR text field
        self.ocr_text = tk.Text(parent, height=10, width=80, bg='white', fg='black', font=("Arial", 12))
        self.ocr_text.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        
        # Bottom frame with buttons
        bottom_frame = ttk.Frame(parent)
        bottom_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.install_tesseract_btn = ttk.Button(bottom_frame, text="Download Tesseract", command=self.install_tesseract)
        self.install_tesseract_btn.pack(side=tk.LEFT, padx=5)
        
        self.btn_copy_ocr = ttk.Button(bottom_frame, text="Copy Text", command=self.copy_ocr_text)
        self.btn_copy_ocr.pack(side=tk.LEFT, padx=5)
        
        self.btn_send_to_main = ttk.Button(bottom_frame, text="Send to Main Window", command=self.send_to_main_window)
        self.btn_send_to_main.pack(side=tk.LEFT, padx=5)
        
        self.lang_label = ttk.Label(bottom_frame, text="Language: -")
        self.lang_label.pack(side=tk.RIGHT, padx=5)
        
        # Tesseract status
        tesseract_status = "Installed" if os.path.exists(self.get_tesseract_path()) else "Not Found"
        self.tesseract_label = ttk.Label(bottom_frame, text=f"Tesseract: {tesseract_status}")
        self.tesseract_label.pack(side=tk.RIGHT, padx=5)
        
        # Хранение изображения
        self.current_image = None
        self.current_pil_image = None

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Текстовые файлы", "*.txt *.md *.json")])
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.text_input.delete('1.0', tk.END)
                    self.text_input.insert(tk.END, f.read())
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть файл: {str(e)}")

    def load_from_clipboard(self):
        try:
            clipboard_text = pyperclip.paste()
            self.text_input.delete('1.0', tk.END)
            self.text_input.insert(tk.END, clipboard_text)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить текст из буфера: {str(e)}")

    def copy_to_clipboard(self):
        try:
            result_text = self.result_text.get('1.0', tk.END)
            if result_text.strip():
                pyperclip.copy(result_text)
                messagebox.showinfo("Успех", "Текст скопирован в буфер обмена")
            else:
                messagebox.showwarning("Предупреждение", "Нет текста для копирования")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось скопировать в буфер обмена: {str(e)}")

    def clean_text(self, text):
        # Удаление ведущих markdown-символов и пробелов
        text = re.sub(r'^[\*\#\-\+]+\s*', '', text, flags=re.MULTILINE)

        # Упрощение длинных дефисов
        text = re.sub(r'—+', '—', text)

        # Полное удаление всех звездочек (включая ** и *)
        text = re.sub(r'\*+', '', text)

        # Удаление лишних пробелов в начале строк
        text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)

        # Обработка форматирования
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Удаление ** (жирный текст в markdown)
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Удаление * (курсив в markdown)
        text = re.sub(r'`(.*?)`', r'\1', text)        # Удаление ` (код в markdown)
        text = re.sub(r'__(.*?)__', r'\1', text)      # Удаление __ (подчеркивание в markdown)
        text = re.sub(r'~~(.*?)~~', r'\1', text)      # Удаление ~~ (зачеркивание в markdown)

        # Обработка списков
        text = re.sub(r'^(\-|\*)\s+', '• ', text, flags=re.MULTILINE)
        text = re.sub(r'^(\d+\.)\s+', r'\1 ', text, flags=re.MULTILINE)

        # Удаление лишних пробелов в начале и конце строк
        text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s+$', '', text, flags=re.MULTILINE)

        # Удаление лишних пустых строк (3+)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Удаление множественных пробелов между словами
        text = re.sub(r'[ ]{2,}', ' ', text)

        # Восстановление переносов после маркированных списков
        text = re.sub(r'(\n• )', r'\n\n• ', text)  # Добавляем перенос перед маркерами
        text = re.sub(r'(\n\d+\.\s)', r'\n\n\1', text)  # Добавляем перенос перед нумерованными пунктами

        # Обработка ссылок Markdown
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', text)  # Заменить [текст](ссылка) на текст

        # Обработка цитат
        text = re.sub(r'^>\s*(.*?)$', r'\1', text, flags=re.MULTILINE)

        # Обработка заголовков
        text = re.sub(r'^#{1,6}\s*(.*?)$', r'\1', text, flags=re.MULTILINE)

        return text.strip()

    def process_tables(self, text):
        tables = []
        lines = text.split('\n')
        current_table = []
        in_table = False

        for line in lines:
            stripped = line.strip()
            if stripped and stripped.startswith('|'):
                if not in_table:
                    in_table = True
                current_table.append(line.strip('|').strip())
            else:
                if in_table:
                    if current_table:
                        tables.append(current_table)
                        current_table = []
                    in_table = False
        if current_table:
            tables.append(current_table)

        processed_tables = []
        for table in tables:
            if not table:  # Проверка на пустые таблицы
                continue
                
            rows = []
            headers = []
            separator_found = False
            for idx, row in enumerate(table):
                cells = [cell.strip() for cell in row.split('|')]
                if idx == 0:
                    headers = cells
                else:
                    if all(re.match(r'^[-\s]*$', cell) for cell in cells):
                        separator_found = True
                    else:
                        if separator_found:
                            rows.append(cells)
                        else:
                            if idx > 0:
                                rows.append(cells)
            processed_tables.append((headers, rows))
        return processed_tables

    def save_to_word(self):
        raw_text = self.text_input.get('1.0', tk.END)
        processed_text = self.clean_text(raw_text)
        tables = self.process_tables(processed_text)

        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'

        for para in processed_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph()
                p.add_run(para).bold = False
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for header, rows in tables:
            if not header:  # Проверка на пустые таблицы
                continue
            table = doc.add_table(rows=1, cols=len(header))
            hdr_cells = table.rows[0].cells
            for i in range(len(header)):
                hdr_cells[i].text = header[i]
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for row in rows:
                row_cells = table.add_row().cells
                for i in range(len(row)):
                    row_cells[i].text = row[i]
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            table.style = 'Table Grid'
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = 1

    def find_table_positions(self, text):
        """Находит позиции таблиц в тексте"""
        table_positions = []
        lines = text.split('\n')
        start_idx = None
        
        for i, line in enumerate(lines):
            if line.strip() and line.strip().startswith('|'):
                if start_idx is None:
                    start_idx = i
            elif start_idx is not None:
                table_positions.append((start_idx, i - 1))
                start_idx = None
                
        if start_idx is not None:  # Если таблица в конце текста
            table_positions.append((start_idx, len(lines) - 1))
            
        return table_positions

    def save_to_word(self):
        raw_text = self.text_input.get('1.0', tk.END)
        processed_text = self.clean_text(raw_text)
        
        # Найдем таблицы и их позиции в тексте
        table_positions = self.find_table_positions(processed_text)
        tables = self.process_tables(processed_text)
        
        doc = Document()
        
        # Установка стиля по умолчанию: Times New Roman, 14pt, полуторный интервал
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)
        
        # Настройка полей страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
        
        # Разбиваем текст на параграфы с учетом позиций таблиц
        lines = processed_text.split('\n')
        current_line = 0
        table_index = 0
        
        for start, end in table_positions:
            # Добавляем текст до таблицы
            if current_line < start:
                text_chunk = '\n'.join(lines[current_line:start])
                for para in text_chunk.split('\n\n'):
                    if para.strip():
                        p = doc.add_paragraph()
                        p.add_run(para).bold = False
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            # Добавляем таблицу, если она есть в списке обработанных таблиц
            if table_index < len(tables):
                header, rows = tables[table_index]
                if header:  # Проверка на пустые таблицы
                    try:
                        table = doc.add_table(rows=1, cols=len(header))
                        table.style = 'Table Grid'
                        table.autofit = True
                        
                        hdr_cells = table.rows[0].cells
                        for i in range(len(header)):
                            hdr_cells[i].text = header[i]
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(14)
                            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        for row in rows:
                            row_cells = table.add_row().cells
                            for i in range(min(len(row), len(row_cells))):
                                row_cells[i].text = row[i]
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(14)
                                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        for row in table.rows:
                            for cell in row.cells:
                                cell.vertical_alignment = 1
                        
                        # Добавим пустой параграф после таблицы для лучшего форматирования
                        doc.add_paragraph()
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Не удалось создать таблицу: {str(e)}")
                
                table_index += 1
            
            current_line = end + 1
        
        # Добавляем оставшийся текст после последней таблицы
        if current_line < len(lines):
            text_chunk = '\n'.join(lines[current_line:])
            for para in text_chunk.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph()
                    p.add_run(para).bold = False
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word Documents", "*.docx")])
        if file_path:
            try:
                doc.save(file_path)
                messagebox.showinfo("Успех", "Документ сохранён")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить документ: {str(e)}")

    def process_text(self):
        raw_text = self.text_input.get('1.0', tk.END).strip()
        processed_text = self.clean_text(raw_text)
        self.result_text.delete('1.0', tk.END)
        self.result_text.insert(tk.END, processed_text)
    
    def load_image(self):
        """Загружает изображение из файла"""
        file_path = filedialog.askopenfilename(filetypes=[
            ("Изображения", "*.png *.jpg *.jpeg *.bmp *.gif")])
        if file_path:
            try:
                image = Image.open(file_path)
                self.current_pil_image = image
                self.display_image(image)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить изображение: {str(e)}")
    
    def load_clipboard_image(self):
        """Загружает изображение из буфера обмена"""
        try:
            # Получаем изображение из буфера обмена
            image = ImageGrab.grabclipboard()
            
            if image is None:
                messagebox.showwarning("Предупреждение", "В буфере обмена нет изображения")
                return
                
            self.current_pil_image = image
            self.display_image(image)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить изображение из буфера: {str(e)}")
    
    def display_image(self, pil_image):
        """Отображает PIL изображение в интерфейсе"""
        if pil_image:
            # Изменяем размер изображения, чтобы оно поместилось в интерфейсе
            width, height = pil_image.size
            max_width = 780
            max_height = 300
            
            # Сохраняем соотношение сторон
            if width > max_width or height > max_height:
                ratio = min(max_width / width, max_height / height)
                new_width = int(width * ratio)
                new_height = int(height * ratio)
                pil_image = pil_image.resize((new_width, new_height), Image.LANCZOS)
            
            # Конвертируем в формат, который понимает tkinter
            self.current_image = ImageTk.PhotoImage(pil_image)
            
            # Обновляем изображение в метке
            self.image_label.config(image=self.current_image)
    
    def process_image(self):
        """Распознает текст на изображении с помощью OCR"""
        if not self.current_pil_image:
            messagebox.showwarning("Предупреждение", "Сначала загрузите изображение")
            return
        
        try:
            # Проверка наличия Tesseract
            if not os.path.exists(self.get_tesseract_path()):
                messagebox.showinfo("Информация", "Tesseract OCR не найден. Попробуем встроенный метод распознавания...")
                # Реализовать альтернативный метод распознавания или предложить пользователю установить Tesseract
                text = self.fallback_ocr()
            else:
                # Извлекаем текст из изображения с помощью Tesseract
                text = pytesseract.image_to_string(self.current_pil_image, lang='rus+eng')
                # Добавляем дополнительные переносы строк для лучшей читаемости
                text = re.sub(r'([.!?])\s*([A-ZА-Я])', r'\1\n\n\2', text)  # Добавляем пустую строку после конца предложения
                text = re.sub(r'\n{3,}', '\n\n', text)  # Убираем лишние пустые строки
            
            if not text.strip():
                messagebox.showinfo("Информация", "Не удалось распознать текст на изображении.")
                return
                
            # Определяем язык текста с учетом возможного смешения языков
            lang = "неизвестный"
            try:
                if text.strip():
                    from langdetect import detect_langs
                    # Получаем список вероятностей для каждого языка
                    langs = detect_langs(text)
                    # Сортируем по вероятности
                    langs.sort(key=lambda x: x.prob, reverse=True)
                    
                    # Если есть хотя бы один определенный язык
                    if langs:
                        primary_lang = langs[0]
                        # Если основной язык русский
                        if primary_lang.lang == 'ru':
                            if len(langs) > 1 and langs[1].lang == 'en' and langs[1].prob > 0.2:
                                lang = "русский + английский"
                            else:
                                lang = "русский"
                        # Если основной язык английский
                        elif primary_lang.lang == 'en':
                            if len(langs) > 1 and langs[1].lang == 'ru' and langs[1].prob > 0.2:
                                lang = "русский + английский"
                            else:
                                lang = "английский"
                        else:
                            lang = f"другой ({primary_lang.lang})"
            except LangDetectException:
                lang = "не определен"
            
            # Обновляем метку с языком
            self.lang_label.config(text=f"Язык: {lang}")
            
            # Отображаем распознанный текст
            self.ocr_text.delete('1.0', tk.END)
            self.ocr_text.insert(tk.END, text)
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка распознавания текста: {str(e)}")

    def fallback_ocr(self):
        """Резервный метод распознавания текста для случаев, когда Tesseract недоступен"""
        # Здесь мы просто показываем сообщение, что для распознавания текста нужно установить Tesseract
        messagebox.showinfo("Информация", 
                          "Для распознавания текста необходимо установить Tesseract OCR.\n\n"
                          "macOS: brew install tesseract tesseract-lang\n"
                          "Windows: скачайте установщик с github.com/UB-Mannheim/tesseract/wiki")
        return ""
    
    def copy_ocr_text(self):
        """Копирует распознанный текст в буфер обмена"""
        text = self.ocr_text.get('1.0', tk.END)
        if text.strip():
            pyperclip.copy(text)
            messagebox.showinfo("Успех", "Текст скопирован в буфер обмена")
        else:
            messagebox.showwarning("Предупреждение", "Нет текста для копирования")
    
    def send_to_main_window(self):
        """Отправляет распознанный текст в основное окно для обработки"""
        text = self.ocr_text.get('1.0', tk.END).strip()
        if text:
            self.text_input.delete('1.0', tk.END)
            self.text_input.insert(tk.END, text)
            messagebox.showinfo("Успех", "Текст отправлен в основное окно")
        else:
            messagebox.showwarning("Предупреждение", "Нет текста для отправки")

if __name__ == "__main__":
    root = tk.Tk()
    app = GPTTextCleaner(root)
    root.mainloop()
